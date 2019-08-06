import xlrd
import string
import re
from nltk.stem import WordNetLemmatizer
from nltk.corpus import wordnet as wn
from docx import Document
from docx.shared import RGBColor

names = set()
with open('names.txt', 'r') as f:
    for line in f:
        names.add(line.rstrip())

patchWords = set()
with open('patch.txt', 'r') as f:
    for line in f:
        patchWords.add(line.rstrip())

commonWords = set()
with open('commonWords.txt', 'r') as f:
    for line in f:
        commonWords.add(line.rstrip())

#workbook = xlrd.open_workbook('fakeSessionNotes_dummyFile.xlsx')
#workbook = xlrd.open_workbook('AllData_sample.xlsx')
workbook = xlrd.open_workbook('AllData.xlsx')
worksheet = workbook.sheet_by_index(0)
nrows = worksheet.nrows
ncols = worksheet.ncols

document = Document()
document.add_heading('Redacted Notes', level=1)
p = document.add_paragraph()
p.add_run('This document contains session notes with names redacted by a program. The words identified as names by the program are denoted by ')
p.add_run('[NAME]').font.color.rgb = RGBColor(255,0,0)
p.add_run(' tokens.')
p = document.add_paragraph()
p.add_run('The program also suggests words which it thinks could be names. These words are not redacted but indicated in ')
p.add_run('blue').font.color.rgb = RGBColor(0,0,255)
p.add_run(' color.')
table = document.add_table(rows=40000, cols=1)

namesList = []
namesCandidates = []
# Bad code, hardcoding the column in which the data is present
# But this will do for now!
notes = worksheet.col(0, start_rowx=0)
lmtr = WordNetLemmatizer()
note_count = 0
for note, row in zip(notes, table.rows):
    note_count += 1
    print(note_count)
#    print('Original Note:\n', note.value)
#    print(note.value)
#    row = table.add_row()
    p = row.cells[0].add_paragraph()
    wordsInNote = note.value.split()
#    print(wordsInNote)
    namesRemoved = []
    for word in wordsInNote:
        # Only consider the words in which only the first letter is capitalized while checking the list of common words. This is because
        # these are the only canditates for being names.
        m = re.search(r'[A-Z][a-z]+', word)
        if m:
        # group(0) returns the entire match.
            stemWord = m.group(0)
        else:
            stemWord = ''
#        print('word', word)
#        print('stem', stemWord)
        synonyms = []
#        antonyms = []
        for syn in wn.synsets(stemWord):
            for l in syn.lemmas():
                synonyms.append(l.name().lower())
#                if l.antonyms():
#                    antonyms.append(l.antonyms()[0].name())

#       print(set(synonyms))
        redact = False
        suggest = False
        # liberal algorithm : identifies namesCandidates, used to suggest redactions.
        if (
            stemWord and
            stemWord.lower() in names or
            (
                stemWord.lower() not in commonWords and
                stemWord.lower() not in patchWords and
                not set(synonyms) & commonWords and
                lmtr.lemmatize(stemWord.lower(), pos='v') == stemWord.lower() and
                # Ignores common nouns which occur as plurals. The singular form is likely included in the list of common words.
                len(lmtr.lemmatize(stemWord.lower(), pos='n')) == len(stemWord)
            )
        ):
            namesCandidates.append(stemWord)
            suggest = True

# conservative algorithm : actually redacts the words
        if (
            stemWord and
            stemWord.lower() in names or
            (
                stemWord.lower() not in commonWords and
                stemWord.lower() not in patchWords and
                not set(synonyms) & commonWords and
#                not set(antonyms) & h and
#            (not wn.synsets(stemWord) or wn.synsets(stemWord)[0].pos() not in ('r','v')) and
#                (not wn.synsets(stemWord) or all([synset.pos() == 'n' for synset in wn.synsets(stemWord)])) and
                (not wn.synsets(stemWord) or all([synset.pos() == 'n' for synset in wn.synsets(stemWord)])) and
                lmtr.lemmatize(stemWord.lower(), pos='v') == stemWord.lower() and
                # Ignores common nouns which occur as plurals. The singular form is likely included in the list of common words.
                len(lmtr.lemmatize(stemWord.lower(), pos='n')) == len(stemWord)
            )
        ):
            namesList.append(stemWord)
            redact = True

        if redact:
            p.add_run('[NAME] ').font.color.rgb = RGBColor(255,0,0)
        elif suggest:
            p.add_run(word + ' ').font.color.rgb = RGBColor(0,0,255)
        else:
            p.add_run(word + ' ')

#document.save('AllDataSample_redacted.docx')
document.save('AllData_redacted.docx')
print('total number of words labeled as names : ', len(namesList))
print('total number of distinct words labeled as names : ', len(set(namesList)))
from collections import Counter
print(Counter(namesList))
print('total number of words suggested as names : ', len(namesCandidates))
print('total number of distinct words suggested as names : ', len(set(namesCandidates)))
print(Counter(namesCandidates))



